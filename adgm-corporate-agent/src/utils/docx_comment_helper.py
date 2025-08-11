# src/utils/docx_comment_helper.py
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.packuri import PackURI
from xml.etree import ElementTree as ET

def _get_or_add_comments_part(doc):
    part = doc.part
    # try to get existing comments part
    rel = None
    for r in part.rels:
        # check relation type string
        try:
            if part.rels[r].reltype == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments":
                rel = part.rels[r]
                break
        except Exception:
            continue
    if rel:
        comments_part = rel.target_part
        comments_xml = ET.fromstring(comments_part._blob)
        return comments_xml, comments_part
    # create comments part
    comments_part = part._package._add_part(
        part._package._next_partname(PackURI("/word/comments.xml")),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml",
        b'<?xml version="1.0" encoding="UTF-8"?><w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"></w:comments>'
    )
    comments_xml = ET.fromstring(comments_part._blob)
    # create relationship from document to comments part
    part.relate_to(comments_part.partname, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments")
    return comments_xml, comments_part

def _save_comments_part(comments_xml, comments_part):
    comments_part._blob = ET.tostring(comments_xml, encoding='utf-8', xml_declaration=True)

def _next_comment_id(comments_xml):
    ids = []
    for c in comments_xml.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comment'):
        cid = c.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
        if cid is not None:
            try:
                ids.append(int(cid))
            except:
                pass
    return max(ids)+1 if ids else 0

def add_comment_to_paragraph(doc: Document, paragraph, comment_text: str, author: str="Reviewer", initials: str="RV"):
    comments_xml, comments_part = _get_or_add_comments_part(doc)
    cid = _next_comment_id(comments_xml)
    # create comment element
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    comment_el = ET.Element('{%s}comment' % ns)
    comment_el.set('{%s}id' % ns, str(cid))
    comment_el.set('{%s}author' % ns, author)
    comment_el.set('{%s}initials' % ns, initials)
    p = ET.Element('{%s}p' % ns)
    r = ET.SubElement(p, '{%s}r' % ns)
    t = ET.SubElement(r, '{%s}t' % ns)
    t.text = comment_text
    comment_el.append(p)
    comments_xml.append(comment_el)
    _save_comments_part(comments_xml, comments_part)
    # insert commentRangeStart, commentRangeEnd and commentReference in paragraph xml
    p_elm = paragraph._p
    start = OxmlElement('w:commentRangeStart')
    start.set(qn('w:id'), str(cid))
    end = OxmlElement('w:commentRangeEnd')
    end.set(qn('w:id'), str(cid))
    # create reference run
    ref_r = OxmlElement('w:r')
    cref = OxmlElement('w:commentReference')
    cref.set(qn('w:id'), str(cid))
    ref_r.append(cref)
    # insert start at beginning and end+ref at end
    p_elm.insert(0, start)
    p_elm.append(end)
    p_elm.append(ref_r)

def annotate_docx_with_comments(input_docx_path: str, output_docx_path: str, issues: list):
    doc = Document(input_docx_path)
    for issue in issues:
        issue_text = issue.get('issue', '')
        suggestion = issue.get('suggestion', '')
        citation = issue.get('citation')
        comment_body = issue_text
        if suggestion:
            comment_body += "\nSuggestion: " + suggestion
        if citation:
            if isinstance(citation, dict) and 'llm_summary' in citation:
                csum = citation['llm_summary']
                if isinstance(csum, dict):
                    comment_body += "\nCitation: " + csum.get('citation','')
                    if csum.get('excerpt'):
                        comment_body += "\nExcerpt: " + csum.get('excerpt')
            else:
                comment_body += "\nCitation: " + str(citation)
        # match paragraph containing keyword
        keywords = [w for w in issue_text.split() if len(w)>3][:6]
        found = False
        for p in doc.paragraphs:
            ptext = p.text or ''
            for kw in keywords:
                if kw.lower() in ptext.lower():
                    add_comment_to_paragraph(doc, p, comment_body)
                    found = True
                    break
            if found:
                break
        if not found:
            last_p = doc.paragraphs[-1]
            add_comment_to_paragraph(doc, last_p, comment_body)
    doc.save(output_docx_path)
